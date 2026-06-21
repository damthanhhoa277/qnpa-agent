# encoding: utf-8
"""
HỌP BAN ĐIỀU HÀNH QNPA — HÀNG TUẦN
Thành phần: Hoa Chủ Tịch · Minh TGĐ · Đào Cố Vấn · Mai Trợ Lý (Thư Ký)
Chủ đề: Mục tiêu & Kết quả hệ thống — KHÔNG bàn thực thi chi tiết
Lịch: Thứ Hai 6:30 — brief chuẩn bị · Sau họp — xuất biên bản Word
"""
import os, sys, json, requests
sys.stdout.reconfigure(encoding="utf-8")
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TELEGRAM_TOKEN = "8608503050:AAEax-HkoZWrWL2QVXP5DXJYhk1FShWd7Zw"
CHAT_COACHING  = "-5017089871"   # Coaching CAS
CHAT_PHONGBAN  = "-5436962310"   # QNPA NỘI BỘ PHÒNG BAN — gửi biên bản họp
VN             = timezone(timedelta(hours=7))

BAO_CAO_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bao_cao_hop")
os.makedirs(BAO_CAO_DIR, exist_ok=True)

# ─────────────────────────────────────────────────────────────
# TELEGRAM
# ─────────────────────────────────────────────────────────────
def tg(chat_id, text):
    try:
        requests.post(
            f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage",
            json={"chat_id": chat_id, "text": text, "parse_mode": "HTML"},
            timeout=10
        )
    except Exception as e:
        print(f"TG loi: {e}")

def tg_file(chat_id, path, caption=""):
    try:
        with open(path, "rb") as f:
            requests.post(
                f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendDocument",
                data={"chat_id": chat_id, "caption": caption},
                files={"document": f},
                timeout=30
            )
    except Exception as e:
        print(f"TG file loi: {e}")

# ─────────────────────────────────────────────────────────────
# BRIEF CHUẨN BỊ — GỬI 6:30 THỨ 2
# ─────────────────────────────────────────────────────────────
def gui_brief_chuan_bi():
    now = datetime.now(VN)
    ngay_str = now.strftime("%d/%m/%Y")
    tuan = _tinh_tuan()

    tg(CHAT_COACHING,
        f"🏛️ <b>HỌP BAN ĐIỀU HÀNH QNPA</b>\n"
        f"📅 Thứ Hai, {ngay_str} · Tuần {tuan}/2026\n"
        f"━━━━━━━━━━━━━━━━━━━━\n\n"
        f"👑 <b>Thành phần:</b>\n"
        f"• Hoa Chủ Tịch (Chủ trì)\n"
        f"• Minh TGĐ (Điều phối)\n"
        f"• Đào Cố Vấn (Phản biện)\n"
        f"• Mai Trợ Lý (Thư ký)\n\n"
        f"🎯 <b>Nguyên tắc cuộc họp:</b>\n"
        f"• Chỉ bàn: Mục tiêu · Kết quả · Rủi ro · Quyết định chiến lược\n"
        f"• KHÔNG bàn: Chi tiết thực thi (CEO làm việc riêng với các phòng)\n"
        f"• Thời lượng: 30-45 phút tối đa\n\n"
        f"📋 <b>AGENDA HÔM NAY:</b>\n"
        f"1️⃣ Kết quả WIG tuần {tuan-1} — đạt/không đạt?\n"
        f"2️⃣ 5 KPI sống chết — xu hướng tốt hay xấu?\n"
        f"3️⃣ Rủi ro nào cần chị Hoa quyết định?\n"
        f"4️⃣ Mục tiêu WIG tuần {tuan} — điều chỉnh nếu cần\n"
        f"5️⃣ Nguồn lực — cần bổ sung gì?\n\n"
        f"💡 <i>Bắt đầu khi chị Hoa sẵn sàng. Minh báo cáo trước.</i>"
    )

    tg(CHAT_COACHING,
        f"📊 <b>SCOREBOARD — TUẦN {tuan-1}</b>\n"
        f"(Lan COO cập nhật số thực vào đây)\n\n"
        f"📥 Inbox/ngày trung bình: ___ (target ≥30)\n"
        f"📞 Lead SĐT/tuần: ___ (target ≥35)\n"
        f"🎓 Học thử/tuần: ___ (target ≥10)\n"
        f"💰 Chốt/tuần: ___ (target ≥5)\n"
        f"💵 CPL trung bình: ___k (target <120k)\n"
        f"📈 Doanh thu tuần: ___tr\n\n"
        f"🎯 WIG tháng: ___ HV mới · Tiến độ: ___/15\n\n"
        f"<i>→ Gõ số thực vào đây trước khi họp</i>"
    )

# ─────────────────────────────────────────────────────────────
# TẠO BIÊN BẢN HỌP — WORD DOC
# ─────────────────────────────────────────────────────────────
def _tinh_tuan():
    return datetime.now(VN).isocalendar()[1]

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _add_heading(doc, text, level=1, color="1B3A6B"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def _add_table_row(table, label, value, bg_label="EBF3FB", bg_value="FFFFFF"):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    _set_cell_bg(row.cells[0], bg_label)
    _set_cell_bg(row.cells[1], bg_value)
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
    return row

def tao_bien_ban(data: dict) -> str:
    """
    data = {
        "tuan": int,
        "ngay": "DD/MM/YYYY",
        "chu_tri": "Hoa Chủ Tịch",
        "ket_qua_wig": "...",
        "kpi": {
            "inbox": "...", "lead": "...", "hoc_thu": "...",
            "chot": "...", "cpl": "...", "doanh_thu": "..."
        },
        "rui_ro": ["rui ro 1", "rui ro 2"],
        "quyet_dinh": ["quyet dinh 1", "quyet dinh 2"],
        "muc_tieu_tuan_sau": "...",
        "nhan_su_bo_sung": "...",
        "y_kien_co_van": "...",
        "bien_ban_vien": "Mai Trợ Lý",
        "ghi_chu": ""
    }
    """
    now = datetime.now(VN)
    ngay = data.get("ngay", now.strftime("%d/%m/%Y"))
    tuan = data.get("tuan", _tinh_tuan())
    ten_file = f"BienBan_HopBanDieuHanh_Tuan{tuan:02d}_{now.strftime('%Y%m%d')}.docx"
    duong_dan = os.path.join(BAO_CAO_DIR, ten_file)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    # ── TIÊU ĐỀ ──────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("HỌC VIỆN PICKLEBALL QNPA")
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(f"BIÊN BẢN HỌP BAN ĐIỀU HÀNH — TUẦN {tuan}/2026")
    r2.bold = True; r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    doc.add_paragraph()

    # ── THÔNG TIN CUỘC HỌP ───────────────────────────────────
    _add_heading(doc, "I. THÔNG TIN CUỘC HỌP")
    info_tbl = doc.add_table(rows=0, cols=2)
    info_tbl.style = "Table Grid"
    info_tbl.columns[0].width = Inches(2.2)
    info_tbl.columns[1].width = Inches(4.3)
    _add_table_row(info_tbl, "Ngày họp", ngay)
    _add_table_row(info_tbl, "Tuần", f"Tuần {tuan}/2026")
    _add_table_row(info_tbl, "Chủ trì", data.get("chu_tri", "Hoa Chủ Tịch — Đàm Thị Thanh Hoa"))
    _add_table_row(info_tbl, "Thư ký", data.get("bien_ban_vien", "Mai Trợ Lý"))
    _add_table_row(info_tbl, "Thành phần", "Hoa Chủ Tịch · Minh TGĐ · Đào Cố Vấn · Mai Trợ Lý")
    _add_table_row(info_tbl, "Hình thức", "Họp nội bộ Ban Lãnh Đạo")

    # ── KẾT QUẢ WIG ─────────────────────────────────────────
    _add_heading(doc, "II. KẾT QUẢ WIG — TUẦN TRƯỚC")
    doc.add_paragraph(data.get("ket_qua_wig", "[Chưa điền]"))

    # ── 5 KPI ────────────────────────────────────────────────
    _add_heading(doc, "III. BÁO CÁO 5 KPI SỐNG CHẾT")
    kpi = data.get("kpi", {})
    kpi_tbl = doc.add_table(rows=0, cols=3)
    kpi_tbl.style = "Table Grid"

    # Header
    hdr = kpi_tbl.add_row()
    for i, txt in enumerate(["KPI", "Thực tế", "Target"]):
        hdr.cells[i].text = txt
        _set_cell_bg(hdr.cells[i], "1B3A6B")
        for para in hdr.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)

    kpi_rows = [
        ("📥 Inbox/ngày trung bình",      kpi.get("inbox", "___"),      "≥ 30"),
        ("📞 Lead SĐT/tuần",              kpi.get("lead", "___"),       "≥ 35"),
        ("🎓 Học thử/tuần",               kpi.get("hoc_thu", "___"),    "≥ 10"),
        ("💰 Chốt/tuần",                  kpi.get("chot", "___"),       "≥ 5"),
        ("💵 CPL trung bình",             kpi.get("cpl", "___"),        "< 120k"),
        ("📈 Doanh thu tuần",             kpi.get("doanh_thu", "___"),  "___"),
    ]
    for label, thuc_te, target in kpi_rows:
        r = kpi_tbl.add_row()
        r.cells[0].text = label
        r.cells[1].text = str(thuc_te)
        r.cells[2].text = target
        _set_cell_bg(r.cells[0], "EBF3FB")
        for cell in r.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

    # ── RỦI RO ───────────────────────────────────────────────
    _add_heading(doc, "IV. RỦI RO & CẢNH BÁO (Đào Cố Vấn)")
    rui_ro = data.get("rui_ro", [])
    if rui_ro:
        for i, rr in enumerate(rui_ro, 1):
            doc.add_paragraph(f"{i}. {rr}", style="List Number")
    else:
        doc.add_paragraph("(Không có rủi ro được ghi nhận)")

    # ── QUYẾT ĐỊNH ───────────────────────────────────────────
    _add_heading(doc, "V. QUYẾT ĐỊNH CỦA HOA CHỦ TỊCH")
    quyet_dinh = data.get("quyet_dinh", [])
    if quyet_dinh:
        qd_tbl = doc.add_table(rows=0, cols=3)
        qd_tbl.style = "Table Grid"
        hdr2 = qd_tbl.add_row()
        for i, txt in enumerate(["#", "Nội dung quyết định", "Giao cho"]):
            hdr2.cells[i].text = txt
            _set_cell_bg(hdr2.cells[i], "1B3A6B")
            for para in hdr2.cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(11)
        for i, qd in enumerate(quyet_dinh, 1):
            if isinstance(qd, dict):
                noi_dung = qd.get("noi_dung", str(qd))
                giao_cho = qd.get("giao_cho", "Minh TGĐ")
            else:
                noi_dung = str(qd)
                giao_cho = "Minh TGĐ"
            row = qd_tbl.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = noi_dung
            row.cells[2].text = giao_cho
            _set_cell_bg(row.cells[0], "EBF3FB")
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(11)
    else:
        doc.add_paragraph("(Chưa có quyết định cần ghi nhận)")

    # ── MỤC TIÊU TUẦN SAU ────────────────────────────────────
    _add_heading(doc, "VI. MỤC TIÊU WIG TUẦN SAU")
    doc.add_paragraph(data.get("muc_tieu_tuan_sau", "[Chưa điền]"))

    # ── NGUỒN LỰC ────────────────────────────────────────────
    _add_heading(doc, "VII. NGUỒN LỰC CẦN BỔ SUNG")
    doc.add_paragraph(data.get("nhan_su_bo_sung", "Không có yêu cầu bổ sung"))

    # ── Ý KIẾN CỐ VẤN ────────────────────────────────────────
    _add_heading(doc, "VIII. Ý KIẾN CỐ VẤN CHIẾN LƯỢC (Đào Cố Vấn)")
    doc.add_paragraph(data.get("y_kien_co_van", "[Chưa điền]"))

    # ── GHI CHÚ ──────────────────────────────────────────────
    ghi_chu = data.get("ghi_chu", "")
    if ghi_chu:
        _add_heading(doc, "IX. GHI CHÚ THÊM")
        doc.add_paragraph(ghi_chu)

    # ── CHỮ KÝ ───────────────────────────────────────────────
    doc.add_paragraph()
    _add_heading(doc, "XÁC NHẬN", level=2)
    ky_tbl = doc.add_table(rows=3, cols=2)
    ky_tbl.style = "Table Grid"
    for i, (ten, chuc) in enumerate([
        ("Đàm Thị Thanh Hoa", "Chủ Tịch — Người phê duyệt"),
        ("Minh TGĐ / AI", "CEO — Người thực thi"),
    ]):
        ky_tbl.rows[0].cells[i].text = chuc
        _set_cell_bg(ky_tbl.rows[0].cells[i], "EBF3FB")
        ky_tbl.rows[1].cells[i].text = "\n\n\n"
        ky_tbl.rows[2].cells[i].text = ten
        for cell in ky_tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10); run.bold = True
    for cell in ky_tbl.rows[2].cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10); run.bold = True

    # Footer note
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run(
        f"Biên bản tạo tự động bởi Mai Trợ Lý — QNPA · {now.strftime('%d/%m/%Y %H:%M')}"
    )
    r_footer.italic = True; r_footer.font.size = Pt(9)
    r_footer.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(duong_dan)
    sys.stdout.buffer.write(f"Da luu: {duong_dan}\n".encode("utf-8", errors="replace"))
    return duong_dan

# ─────────────────────────────────────────────────────────────
# GỬI BIÊN BẢN QUA TELEGRAM
# ─────────────────────────────────────────────────────────────
def gui_bien_ban(duong_dan: str, tuan: int):
    caption = (
        f"📋 BIÊN BẢN HỌP BAN ĐIỀU HÀNH — TUẦN {tuan}/2026\n"
        f"Tạo bởi Mai Trợ Lý · {datetime.now(VN).strftime('%d/%m/%Y %H:%M')}\n"
        f"📂 Lưu tại: bao_cao_hop/"
    )
    tg_file(CHAT_PHONGBAN, duong_dan, caption)
    tg(CHAT_PHONGBAN,
        f"✅ <b>Biên bản Họp Ban Điều Hành đã gửi!</b>\n\n"
        f"Toàn bộ team xem quyết định và mục tiêu tuần này nhé.\n"
        f"📁 Bản gốc lưu tại: <code>bao_cao_hop/</code>"
    )

# ─────────────────────────────────────────────────────────────
# DEMO — CHẠY THỬ VỚI DỮ LIỆU MẪU
# ─────────────────────────────────────────────────────────────
def demo_bien_ban():
    data = {
        "tuan": _tinh_tuan(),
        "ngay": datetime.now(VN).strftime("%d/%m/%Y"),
        "ket_qua_wig": (
            "WIG tháng 7: 15 học viên mới. Tuần này: 3 HV đăng ký (cần ≥5/tuần để đạt mục tiêu).\n"
            "Tiến độ: 3/15 HV — đang chậm so với kế hoạch. Cần tăng tốc từ tuần tới."
        ),
        "kpi": {
            "inbox": "8/ngày 🔴",
            "lead": "12/tuần 🟡",
            "hoc_thu": "4/tuần 🔴",
            "chot": "3/tuần 🟡",
            "cpl": "195k 🔴",
            "doanh_thu": "18tr"
        },
        "rui_ro": [
            "CPL tăng 62% so tháng trước — creative Ads mệt, cần thay gấp trong tuần này",
            "Inbox quá thấp (8/ngày vs target 30) — traffic chưa đủ để đạt WIG",
            "Hải Marketing chưa có Ads mới sau 2 tuần — rủi ro mất momentum tháng 7",
        ],
        "quyet_dinh": [
            {"noi_dung": "Hải Marketing phải có 3 mẫu Ads mới trước thứ 4 (25/6)", "giao_cho": "Hải Marketing"},
            {"noi_dung": "Tăng budget test lên 600k/ngày (3 mẫu × 200k) tuần tới", "giao_cho": "Minh TGĐ"},
            {"noi_dung": "Hùng lấy đủ 5 testimonial từ HV tháng 6 trước thứ 5 (26/6)", "giao_cho": "Hùng Doanh Số"},
        ],
        "muc_tieu_tuan_sau": (
            "• Inbox/ngày: ≥15 (từ Ads mới)\n"
            "• Học thử: ≥6 buổi\n"
            "• Chốt: ≥4 HV\n"
            "• CPL: < 160k (cải thiện từ 195k)"
        ),
        "nhan_su_bo_sung": "Chưa cần bổ sung — ưu tiên tối ưu team hiện tại trước.",
        "y_kien_co_van": (
            "Đào Cố Vấn nhận định: Vấn đề cốt lõi không phải ngân sách mà là CREATIVE mệt.\n"
            "Khuyến nghị: Ưu tiên lấy testimonial video thật (phụ huynh nói camera) — "
            "đây là creative có CPL thấp nhất trong ngành giáo dục."
        ),
        "ghi_chu": "Họp tiếp theo: Thứ Hai tuần sau, 6:30. Mai nhắc chị Hoa Chủ Nhật tối.",
    }
    duong_dan = tao_bien_ban(data)
    return duong_dan, _tinh_tuan()


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    cmd = sys.argv[1] if len(sys.argv) > 1 else ""

    if cmd == "brief":
        # Gửi brief chuẩn bị họp
        print("Gui brief chuan bi hop...")
        gui_brief_chuan_bi()

    elif cmd == "bien_ban":
        # Tạo biên bản mẫu + gửi Telegram
        print("Tao bien ban hop mau...")
        duong_dan, tuan = demo_bien_ban()
        print(f"File: {duong_dan}")
        gui_bien_ban(duong_dan, tuan)

    elif cmd == "tao":
        # Tạo biên bản từ JSON file (dùng sau khi họp thật)
        # Usage: python hop_ban_dieu_hanh.py tao data_hop.json
        if len(sys.argv) < 3:
            print("Usage: python hop_ban_dieu_hanh.py tao <data_hop.json>")
        else:
            with open(sys.argv[2], encoding="utf-8") as f:
                data = json.load(f)
            duong_dan = tao_bien_ban(data)
            tuan = data.get("tuan", _tinh_tuan())
            print(f"Da tao: {duong_dan}")
            gui_bien_ban(duong_dan, tuan)

    else:
        print("Lenh:")
        print("  brief      — Gui brief chuan bi hop (6:30 Thu 2)")
        print("  bien_ban   — Tao bien ban mau + gui Telegram")
        print("  tao <file> — Tao bien ban tu data_hop.json")
